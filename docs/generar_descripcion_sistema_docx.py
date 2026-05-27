"""
Genera descripcion funcional completa del ecosistema Cleexs en Word (.docx).
Requiere: pip install python-docx
Salida por defecto: docs/Descripcion_Funcional_Completa_Cleexs.docx
"""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
out_path = ROOT / "docs" / "Descripcion_Funcional_Completa_Cleexs.docx"

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

title = doc.add_paragraph()
title_run = title.add_run("Cleexs — Descripción funcional completa del proyecto")
title_run.bold = True
title_run.font.size = Pt(20)
title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

doc.add_paragraph(f"Fecha de actualización del documento: {date.today().isoformat()}")
doc.add_paragraph(
    "Alcance: descripción orientada al negocio y al uso funcional del monorepo principal "
    "(apps/web + apps/api + Prisma), más referencia a proyectos satélite en el mismo repositorio."
)


def h1(text: str):
    doc.add_heading(text, level=1)


def h2(text: str):
    doc.add_heading(text, level=2)


def h3(text: str):
    doc.add_heading(text, level=3)


def p(text: str):
    doc.add_paragraph(text)


def bullet(text: str):
    doc.add_paragraph(text, style="List Bullet")


# ---------------------------------------------------------------------------
h1("1. Resumen ejecutivo")
p(
    "Cleexs es una plataforma SaaS multi-tenant para medir, de forma repetible y auditable, "
    "cómo una marca aparece y es recomendada en respuestas de modelos de lenguaje (principalmente "
    "ChatGPT/OpenAI y, en planes superiores, también Gemini). El núcleo del producto es un índice "
    "numerico tipo score (PRIA/Cleexs Score, escala típica 0–100) con evidencia por consulta («prompt»), "
    "por competidor y por corrida («run»), de modo que equipos comerciales y de marketing puedan comparar "
    "resultados entre periodos y versiones de prompts."
)
p(
    "La solución incluye un flujo de diagnóstico público guiado por dominio, dashboards internos de marca "
    "y de plataforma, gestión versionada de prompts, outreach asistido (detección de competidores más "
    "visibles que la marca y enriquecimiento de contactos vía scraping y Hunter.io), así como superficies "
    "de cliente (portal cliente / portal crecimiento) consumiendo permisos y cupos mediante un modelo de "
    "planes y registros de uso."
)
p(
    "Operativamente, el núcleo se despliega como aplicación Next.js en el front y API Fastify sobre "
    "PostgreSQL (Prisma) en backend, con opciones típicas de hosting tipo Vercel + Railway."

)

h1("2. Propuesta de valor y objetivos de negocio")
bullet("Ofrecer una métrica comparativa estable en el tiempo (no un «one-off» cualitativo) sobre recomendaciones de IA.")
bullet(
    "Dar trazabilidad: cada score agregado se puede relacionar con las respuestas crudas, el top de marcas mencionadas y metadatos de calidad."
)
bullet(
    "Conectar métricas con acción comercial mediante outreach: identificar contra quién pierde la marca y acercar evidencia para contactos B2B."
)
bullet(
    "Soportar modelo de planes (límites de marcas, competidores, prompts activos, corridas/mes y automatización) y extensiones mediante overrides administrativos."
)

h1("3. Actores y superficies de la aplicación")
h2("3.1 Visitante / usuario anónimo")
bullet("Acceso al flujo público de diagnóstico, páginas de marketing (landing, planes) y vistas de resultado compartible por slug cuando exista.")
bullet("Sin cuenta obligatoria para iniciar ciertos diagnósticos públicos.")

h2("3.2 Usuario operativo dentro del tenant")
bullet(
    "Gestiona marcas, competidores, aliases, clasificación automática/manual, prompts y corridas desde el dashboard y módulos de configuración/reportes/outreach habituales."
)
bullet("Consume reportes PRIA/Cleexs Score, rankings y comparativas entre periodos/versiones.")

h2("3.3 Usuario portal (credenciales de portal)")
bullet(
    "Usuario con hash de contraseña en tabla User: accede mediante login JWT expuesto por la API (/api/auth/portal/login)."
)
bullet("Vistas específicas bajo rutas tipo portal-cliente y portal-crecimiento orientadas al informe premium y uso de entitlement.")

h2("3.4 Operaciones / administrador interno")
bullet(
    "Panel web bajo rutas protegidas de administración (login admin) que proxifica llamadas firmadas contra la API con secreto ADMIN_API_SECRET."
)
bullet(
    "Funciones típicas: resumen ejecutivo dashboard, alta de cuentas, overrides de entitlement, estadísticas email/Resend, campañas internas."
)

h1("4. Funcionalidades principales")

h2("4.1 Diagnóstico público automatizado")
bullet("Entrada típica: dominio o marca dentro del funnel público (/diagnostico/*, ver resultado, etc.).")
bullet(
    "El backend clasifica contexto industrial y geográfico (TLD, búsqueda y modelo de lengua), infiere país/mercado y deriva prompts de diagnosis alineados a intención de compra."
)
bullet("Detección y validación asistida de competidores (top competidores, resolución de dominio oficial).")
bullet(
    "Ejecución contra OpenAI (y segunda corrida Gemini en nivel «Gold» o equivalente cuando el plan/feature lo habiliten), procesamiento asíncrono o no bloqueante según UX."
)
bullet(
    "Cálculo de PRIA/score marca, métricas por intención, narrativa/análisis textual y resultado persistido con vínculos de uso viral."
)

h2("4.2 Página «score» compartible")
bullet(
    "Resultado público direccionable (ruta tipo /score/{slug}) asociado a PublicDiagnostic/PublicDiagnosticShareVisit para métricas de difusión y seguimiento de visitas compartidas."
)

h2("4.3 Multi-tenant, planes y jerarquía")
bullet(
    "Modelo Tenant con tipos ROOT, AGENCY, DIRECT_CLIENT, AGENCY_CLIENT; caminos tenant_code/tenant_path para jerarquía de reventa y agencias."
)
bullet(
    "Cada Tenant referencia Plan con límites: corridas mensuales, prompts activos, marcas, competidores, meses de retención de datos y flag de automatización periódica."
)
bullet("Estados de tenant: active, suspended, archived.")

h2("4.4 Marcas y competencia")
bullet("CRUD de Brand con dominio, industria, descripción objetivos y clasificación automática (businessType, category, subcategory, geoMarket, sizeSegment) con marca autoDetected/classifier_meta JSON.")
bullet("Aliases por marca para alinear nomenclatura de IA con la marca oficial.")
bullet("Competidor asociado a marca: nombre, opcionalmente dominio para outreach y metadatos de clasificación paralelos.")

h2("4.5 Prompts versionados")
bullet("PromptCategory y Prompt dentro de PromptVersion versionable por tenant: activación, listado y clonado.")
bullet(
    "Permite reproducir cohortes históricas («qué prompts usábamos en Q3») y comparar resultados antes/después."
)

h2("4.6 Corridas (runs), resultados y override manual")
bullet("Run por marca/en periodo con estados pending/running/completed/failed gestionados por API.")
bullet(
    "Cada corrida ejecuta el conjunto de prompts activos: se almacena PromptResult con fragmento de respuesta, top-3 parseado, flags de calidad y score por prompt."
)
bullet("Override manual por resultado para corregir interpretaciones ambiguas y recalcular agregados.")
bullet("PRIAReport agrega métricas por run/marca para dashboards.")
bullet("Programación sugerida vía campo runSchedule en Brand (semanal/quincenal/mensual) pensada para integración con cron externo (n8n).")

h2("4.7 Reportería y dashboards")
bullet("Reportes de ranking, comparación entre versiones de prompts, dashboard de plataforma (actividad, runs, industrias, diagnósticos trazados) y dashboard de marca (score actual, tendencia, último run).")
bullet("Vistas en Next.js bajo /dashboard, /reports y reportes embebidos en portales según plan.")

h2("4.8 Portal cliente y portal crecimiento")
bullet(
    "Conjunto de páginas dedicadas a presentar el informe al cliente final y la experiencia «crecimiento» con secciones premium: suscripción, comparación, competidores, historial, prompts, herramientas, equipo, etc."
)
bullet("El acceso efectivo a generación profunda/visualización está gobernado por entitlements + usage ledger.")

h2("4.9 Entitlements, overrides y libro de uso")
bullet(
    "Acciones modelo: score_view, score_generate, report_deep_generate, report_deep_view, profile_claim — cada una chequeada combinando límites de plan con overrides específicos al tenant/usuario."
)
bullet(
    "UsageLedger registra consumo mensual típico (vistas/generaciones) ante actor user o anónimo (según feature flags de operación)."
)
bullet(
    "EntitlementOverride activable por administración permite cortesías, extensiones temporales y pruebas comerciales sin mutar permanentemente Plan."
)

h2("4.10 Programa de referidos (portal cliente)")
bullet(
    "Por tenant existe referralSlug estable; un tenant puede referenciarse a otro mediante referredByTenantId durante altas/recovery de portal."
)
bullet(
    "Contadores referralCount, referralRewardAt, flags de dismissed upsell permiten automatizar beneficios cuando el ciclo está habilitado en negocio."
)

h2("4.11 Perfil reclamado (profile claim)")
bullet(
    "Entidad ProfileClaim con estados pendiente/aprobado/rechazado para flujos en los que una marca física debe verificar dominio/perfil público dentro de Cleexs."
)

h2("4.12 Outreach — pipeline comercial MVP")
bullet(
    "Tras corrida diagnóstica, el sistema identifica competidores que aparecen mejor posicionados que la marca en conjuntos evidenciados por prompts."
)
bullet(
    "Crea LeadSource con payload de evidencias (pregunta exacta, top3, orden) y permite enriquecer contactos (LeadContact) con emails descubiertos y borradores (LeadEmail)."
)
bullet("Fuentes de enriquecimiento: estrategia Firecrawl map+scrape sobre subconjuntos de URLs más Hunter.io Domain Search.")
bullet(
    "MVP enfocado en visibilización y preparación; el envío masivo automatizado outbound puede estar parcial/no productivo según configuración SMTP/ESP."
)

h2("4.13 Provisionamiento administrativo de cuentas")
bullet(
    "Endpoint protegido /api/admin/provision-account: crea Tenant+User+Brand inicial y password de portal; soporta plan free vs crecimiento, cortesia de entitlement y slug de referidor."
)
bullet(
    "Herramientas auxiliares: randomización de cuenta demo desde corrida reciente, ensure user portal para marca existente (según rutas definidas)."
)
bullet(
    "En CLI local equivalente mediante scripts npm tipo db:provision:account mencionados en mensajes API."
)

h2("4.14 Comunicaciones internas Cleexs (email operativo)")
bullet(
    "Modelo CleexsInternalEmailCampaign: campaña interna parametrizada por semana bucket de score y plantillas ESP opcionales (espTemplateId Resend)."
)
bullet(
    "CleexsInternalEmailSendLog audita cada intento/aprobación SMTP o API outbound (estado externo mergeado con metadata)."
)
bullet("Envíos de prueba administrados por API (/api/admin/email/*) usando credenciales Resend/API key o relays configurados.")

h2("4.15 Webhooks de Resend (engagement métricas)")
bullet(
    "Endpoint /api/webhooks/resend ingiere eventos email.sent/delivered/opened/etc. firmados Svix cuando RESEND_WEBHOOK_SECRET coincide con signing secret de Resend."
)
bullet(
    "Persistence en tabla cleexs_resend_webhook_events; panel admin muestra conteos últimos N días. Sin secreto válido la API rechaza ingestion (ej. HTTP 503) y métricas quedan vacías."
)

h2("4.16 Admin UI — cuentas, overrides y estadísticas globales")
bullet(
    "/admin/(interno)/cuentas: búsqueda operativa sin jergas técnicas, acciones tipo cortesia y visualización rápida de overrides."
)
bullet(
    "/admin/(interno)/email: KPIs configuración SMTP/relay, estadísticas campañas, logs recientes y disparadores de envío/prueba enlazados a estadísticas webhook."
)

h2("4.17 Autenticación portal y sesión Bearer")
bullet("POST /api/auth/portal/login valida bcrypt sobre password_hash y emitir JWT firmado PORTAL_JWT_SECRET.")
bullet("POST /api/auth/portal/me retorna usuario autenticado si Authorization Bearer válido.")

h2("4.18 Cron programado externo")
bullet(
    "Endpoint /api/cron/scheduled-runs pensado para n8n o scheduler externo usando CRON_SECRET en cabeceras; dispara correlación de corrida automática pendiente por plan/marca."
)

h2("4.19 Página Facturas")
bullet(
    "Sección navegable /facturas en el front como placeholder UX: comunica histórico de comprobantes futuro cuando exista billing integrado."
)

h1("5. Integraciones externas")
bullet("OpenAI (Chat Completions) — clasificación, generación textual, corrida prompts.")
bullet("Google Gemini API (tier Gold cuando aplica).")
bullet("Firecrawl v2 — map + scrape + extracción heurística de emails.")
bullet("Hunter.io — domain search de contactos verificados/puntuados.")
bullet("Resend (u otro SMTP relay) — transaccional, campañas internas y webhooks de engagement.")
bullet("Infra: despliegue típico Vercel frontend + Railway API y base Postgres gestionada.")

h1("6. Arquitectura técnica resumida")
h2("6.1 Monorepo núcleo")
bullet("apps/web — Next.js 14/React/TypeScript/Tailwind, consume API REST mediante variables NEXT_PUBLIC_* y rutas servidor para proxy seguro donde aplica.")
bullet("apps/api — Fastify + plugins CORS (whitelist expandible incl. dominios cliente y healthcheck Railway), Helmet y fastify-raw-body para rutas webhook firmadas.")
bullet("packages/shared — utilidades score/parser compartidas front/back.")
bullet("prisma/schema.prisma — fuente única del modelo PostgreSQL.")

h2("6.2 Salud operativa API")
bullet("GET /health — estado proceso y diagnostico opcional integrations.resendWebhookSecretConfigured cuando desplegada versión reciente.")

h2("6.3 Endpoints destacados por prefijo (referencia rápida)")
p("Todos bajo mismo host salvo rutas públicas aisladas; prefijo habitual /api.")
bullet("/api/public/* — ciclo vida diagnóstico público (crear/consultar/resultado/compartición).")
bullet("/api/brands,/api/prompts,/api/runs,/api/reports — núcleo operativo tenant.")
bullet("/api/leads — descubrir leads y crear borradores email.")
bullet("/api/me/* y /usage — consumo entitlement autenticado portal.")
bullet("/api/admin/* — operaciones provisioning, email, overrides, dashboard summary (cabecera x-admin-secret).")
bullet("/api/webhooks/resend — ingestion firmada webhook.")
bullet("/api/auth/portal/* — login JWT cliente.")

h1("7. Modelo de datos — entidades nucleares")
p("Tablas modelo (PostgreSQL via Prisma) — relación función:")
bullet("Tenant — nodo tenancy con plan_id, estado, referidos jerárquicos.")
bullet("Plan — definición límites y precio opcional mensual.")
bullet("User — email único dentro del tenant + rol owner/editor/viewer.")
bullet("Brand — marca medida dentro del tenant.")
bullet("BrandAlias, Competitor — normalización rivalry.")
bullet("PromptVersion → PromptCategory → Prompt — colección ejecutable.")
bullet("Run — instancia período marca; PromptResult línea por línea ejecutada; PRIAReport agregados.")
bullet("PublicDiagnostic + PublicDiagnosticShareVisit — resultado viral.")
bullet("UsageLedger — contabilidad entitlement.")
bullet("TenantBrandAccess — compartución cross-tenant marca si aplica.")
bullet("EntitlementOverride — cortesías administrativas.")
bullet("ProfileClaim — verificación de perfil marca.")
bullet("LeadSource, LeadContact, LeadEmail — pipeline outreach.")
bullet("CleexsInternalEmailCampaign, CleexsInternalEmailSendLog — operación comunicaciones internas.")
bullet("CleexsResendWebhookEvent — auditoría tiempo real eventos ESP.")

h1("8. Proyectos Satélite (mismo repo, otros stacks)")
p(
    "Además del monorepo TypeScript núcleo, el directorio Satelite agrupa herramientas standalone — útiles demo/SEO técnico — habitualmente Backend Python FastAPI + Frontend Next dedicado cada uno."
)

h3("8.1 Project1 — Crawlability Checker")
bullet("Rastrea sitio, SEO básico, robots y bots IA; produce score crawlability.")

h3("8.2 Project2 — Robots & Sitemap Analyzer")
bullet("Analiza robots.txt + sitemap, visibilidad bots búsqueda vs IA, sugerencias de archivos.")

h3("8.3 CleexsTools37 / frontendToni / backendToni")
bullet(
    "Variantes experimentales o demo (panel tipo dashboard de herramientas SEO/IA) no sustituyen al producto principal apps/web+api; sirven prototipos o integraciones separadas."
)

h1("9. Seguridad y control operativo")
bullet("CORS restringido salvo listas blancas y flags de entorno ALLOW_* muy acotadas.")
bullet("Helmet en API.")
bullet(
    "Rutas admin y webhook requieren secretos dedicados (ADMIN_API_SECRET, RESEND_WEBHOOK_SECRET, CRON_SECRET) — nunca commiteados."
)
bullet("Contraseña portal bcrypt; JWT TTL configurable por implementación.")

h1("10. Variables de entorno críticas (sin valores)")
bullet("DATABASE_URL")
bullet("OPENAI_API_KEY / GEMINI u GOOGLE según modo")
bullet("FIRECRAWL_API_KEY, HUNTER_API_KEY")
bullet("FRONTEND_URLS o FRONTEND_URL para CORS")
bullet("ADMIN_API_SECRET, CRON_SECRET, PORTAL_JWT_SECRET")
bullet("Credenciales Resend/SMTP (API key, relay SMTP) + RESEND_WEBHOOK_SECRET + opcional PUBLIC_WEBHOOK_BASE_URL")
bullet("NEXT_PUBLIC_API_URL (front)")

h1("11. Estado de madurez y límites reconocidos")
p("Madurez: MVP-plus operativo sobre flujo núcleo + outreach + portals + automatización programable externamente.")
bullet("Fortalezas: trazabilidad score, modelo datos rico para upsell/reporting, automatización diagnóstico→lead.")
bullet("Límites típicos: dependencia datos web externos sin email visible, colas heavyweight en proceso app (sin worker dedicado forzado del repo madre).")

h1("12. Extensiones naturales próximas")
bullet("Billing nativo Stripe u otro PSP enlazando facturas UI.")
bullet(
    "Workers desacoplados outreach largo."
)
bullet(
    "Más enrichment providers y mejor control reputación outbound."
)

h1("13. Cierre")
p(
    "Cleexs converge medición reputacional IA, evidencia ejecutable para ventas agencia/marca, y modelo SaaS modular con límites reales y controles administrativos. Esta descripción debe leerse como mapa "
    "funcional del codebase principal; proyectos Satelite son satellites de utilidad técnico-comercial paralelos."

)

doc.save(out_path.as_posix())
print(f"Generado: {out_path}")
